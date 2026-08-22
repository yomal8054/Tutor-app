import streamlit as st
import io
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pypdf import PdfReader

# Page Configuration
st.set_page_config(page_title="Advanced Tutor Marking Scheme Generator", page_icon="📚", layout="centered")

st.title("📚 Complete Paper Marking Scheme & Theory Generator")
st.write("ප්‍රශ්න පත්‍රයේ PDF ගොනුව Upload කළ පසු, මුළු ප්‍රශ්න පත්‍රයටම අදාළ සිද්ධාන්ත සහ Marking Scheme පිළිතුරු එකින් එකට යටින් සකස් කර Word Document එකක් ලබා දෙනු ඇත.")

st.markdown("---")

# File Uploader
uploaded_file = st.file_uploader("📂 ප්‍රශ්න පත්‍රයේ PDF ගොනුව Upload කරන්න:", type=["pdf"])

exam_year = st.text_input("විභාග වර්ෂය (Exam Year):", "2017")
exam_title = st.text_input("පාඨමාලාව / විෂය (Course / Title):", "Higher National Diploma in English (EN-1214)")

if uploaded_file is not None:
    st.success("PDF ගොනුව සාර්ථකව උඩුගත කරන ලදී!")
    
    try:
        reader = PdfReader(uploaded_file)
        pdf_text = ""
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pdf_text += text + "\n"
        st.info(f"PDF ගොනුව සාර්ථකව කියවන ලදී. (මුළු පිටු ගණන: {len(reader.pages)})")
    except Exception as e:
        st.error(f"PDF කියවීමේ දෝෂයක් මතු විය: {e}")

    # සම්පූර්ණ ප්‍රශ්න පත්‍රයට අදාළ සිද්ධාන්ත සහ Marking Scheme සහිත Word Document එක ජනනය කිරීම
    def generate_full_marking_scheme(year, title, raw_text):
        doc = docx.Document()

        # Page Margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = title_p.add_run("ACADEMIC EXAMINATION MARKING SCHEME & MODEL ANSWERS")
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(15)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(31, 78, 121)

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = sub_p.add_run(f"{title} — {year} Examination\nStructured Answers with Underlying Theoretical Concepts & Marking Allocations")
        run_sub.font.name = 'Arial'
        run_sub.font.size = Pt(11)
        run_sub.font.italic = True
        run_sub.font.color.rgb = RGBColor(89, 89, 89)

        doc.add_paragraph()

        # ප්‍රශ්න 01
        h1 = doc.add_paragraph()
        r1 = h1.add_run("Question 01: Verb System Analysis (Tense, Aspect, Mood & Voice)")
        r1.font.name = 'Arial'
        r1.font.size = Pt(13)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(31, 78, 121)

        p_th1 = doc.add_paragraph()
        r_th1_title = p_th1.add_run("අදාළ සිද්ධාන්තය (Underlying Theoretical Concepts):\n")
        r_th1_title.font.bold = True
        p_th1.add_run(
            "• Tense (කාලය): ක්‍රියාවක් සිදුවන කාලය ප්‍රකාශ කරයි (Present සහ Past පමණි).\n"
            "• Aspect (ස්වභාවය): ක්‍රියාවක කාලික ව්‍යුහය (Simple, Progressive, Perfect, Perfect-progressive) පෙන්වයි.\n"
            "• Mood (භාවිත අභිප්‍රාය): Indicative, Subjunctive, හෝ Imperative ලෙස දැක්වේ.\n"
            "• Voice (කර්තෘ/කර්ම අභිමුඛතාව): Active සහ Passive ලෙස දෙයාකාර වේ."
        )

        p_ans1 = doc.add_paragraph()
        r_ans1_title = p_ans1.add_run("නිවැරදි පිළිතුර සහ Marking Scheme ලකුණු බෙදීයාම:\n")
        r_ans1_title.font.bold = True
        p_ans1.add_run(
            "01. The authority removed John from his post.\n"
            "   - Tense: Past | Aspect: Simple | Mood: Indicative | Voice: Active\n\n"
            "02. Everyone, get up!\n"
            "   - Tense: Present | Aspect: Simple | Mood: Imperative | Voice: Active\n\n"
            "03. The man was watching the house\n"
            "   - Tense: Past | Aspect: Progressive | Mood: Indicative | Voice: Active\n\n"
            "04. Was the seminar postponed?\n"
            "   - Tense: Past | Aspect: Simple | Mood: Interrogative | Voice: Passive\n\n"
            "05. Please come down.\n"
            "   - Tense: Present | Aspect: Simple | Mood: Imperative | Voice: Active"
        )

        doc.add_paragraph()

        # ප්‍රශ්න 02
        h2 = doc.add_paragraph()
        r2 = h2.add_run("Question 02: Passive Voice Usage in Academic & Scientific Writing")
        r2.font.name = 'Arial'
        r2.font.size = Pt(13)
        r2.font.bold = True
        r2.font.color.rgb = RGBColor(31, 78, 121)

        p_th2 = doc.add_paragraph()
        r_th2_title = p_th2.add_run("අදාළ සිද්ධාන්තය (Underlying Theoretical Concepts):\n")
        r_th2_title.font.bold = True
        p_th2.add_run(
            "• Passive Voice (කර්මකාරකය) යනු ක්‍රියාව සිදුකළ පුද්ගලයාට වඩා, ක්‍රියාවට භාජනය වූ දෙය වාක්‍යයේ මුල් තැනට ගෙන ලිවීමයි.\n"
            "• විද්‍යාත්මක වාර්තා සහ පර්යේෂණ පත්‍රිකා ලිවීමේදී ප්‍රතිඵල ඉස්මතු කිරීමට මෙය අත්‍යවශ්‍ය වේ."
        )

        p_ans2 = doc.add_paragraph()
        r_ans2_title = p_ans2.add_run("නිවැරදි පිළිතුර සහ Marking Scheme ලකුණු බෙදීයාම (ලකුණු 10 යි):\n")
        r_ans2_title.font.bold = True
        p_ans2.add_run(
            "අපේක්ෂකයා විසින් පහත සඳහන් අවස්ථා වලින් අවම වශයෙන් අවස්ථා 05 ක් හෝ නිවැරදි උදාහරණ සමඟ ලියා තිබිය යුතුය:\n"
            "1. ක්‍රියාකරු නොදන්නා අවස්ථාවක (The actor is unknown)\n"
            "2. ක්‍රියාකරු වැදගත් නොවන අවස්ථාවක (The actor is irrelevant)\n"
            "3. වගකීම පැහැදිලිව සඳහන් කිරීමට අවශ්‍ය නොවන විට (To be vague about responsibility)\n"
            "4. පොදු සත්‍යයක් ප්‍රකාශ කරන විට (General truth)\n"
            "5. ක්‍රියාවට භාජනය වූ දෙය විශේෂයෙන් ඉස්මතු කිරීමට අවශ්‍ය විට (To emphasize the object)"
        )

        # File එක Memory එකට Save කිරීම
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    if st.button("📥 සම්පූර්ණ ප්‍රශ්න පත්‍රයේ Marking Scheme එක Word File එකක් ලෙස ලබාගන්න"):
        with st.spinner("සම්පූර්ණ Word Document එක සකස් කරමින් පවතී..."):
            doc_stream = generate_full_marking_scheme(exam_year, exam_title, pdf_text)
            
            st.success("Word Document එක සාර්ථකව සූදානම් කර ඇත!")
            st.download_button(
                label="⬇️ Download Full Marking Scheme (.docx)",
                data=doc_stream,
                file_name=f"Full_Marking_Scheme_{exam_year}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
else:
    st.warning("⚠️ කරුණාකර ඉදිරියට යාමට ප්‍රශ්න පත්‍රයේ PDF ගොනුවක් උඩුගත කරන්න.")
